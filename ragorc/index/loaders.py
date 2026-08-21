"""Loaders: a file becomes a :class:`~ragorc.core.models.Document` with an identity.

Identity is the whole job
-------------------------
Extracting text is the easy half. The half that decides whether a re-ingest costs
nothing or costs the full corpus is how a document is *identified*, and this
package fixes two rules for every loader:

* **The id comes from the source, never from the content.**
  ``document_id(source, tenant_id=...)`` is called without the content argument on
  purpose. Folding the content into the id would give an edited file a brand new
  id on every save, so the store would accumulate a fresh document per edit, the
  old chunks would be orphaned in the index forever, and the checksum comparison
  the pipeline relies on could never match anything.
* **The checksum comes from the content, and it is mandatory.**
  ``Document.checksum = content_hash(content)``. Ingest is idempotent on
  ``(id, checksum)``: :class:`~ragorc.index.pipeline.IngestPipeline` reads the
  stored checksum for the ids it is about to write and skips the ones that did not
  move. That turns a nightly re-ingest of a mostly-static corpus into a metadata
  query, which is the single biggest cost saving available at index time — but only
  if the checksum is present, so no loader here is allowed to omit it.

Multi-record sources (a CSV row, a PDF page, a JSON array element) get a
**fragment** appended to the source — ``report.pdf#page=7``, ``rows.csv#row=41`` —
and the fragment is part of the identity. Row 41 therefore keeps its id when row
42 changes. Where the record carries its own key, that key is preferred over the
ordinal for exactly this reason: identifying a record by its position means
inserting one row at the top re-identifies every row below it, and a re-ingest
then rewrites the entire file's worth of vectors.

Every read happens in a worker thread
-------------------------------------
``Path.read_bytes`` plus UTF-8 decoding of a 10 MB file is tens of milliseconds of
blocking work, and PDF text extraction is far worse. On the event loop that stalls
every in-flight store request and LLM call in the process, so each loader does one
``asyncio.to_thread`` hop covering open, read, decode and parse together — one hop
rather than several, because the thread transition costs more than any single step
inside it.

Heavy parsers are imported inside the function
----------------------------------------------
pymupdf, beautifulsoup4, lxml and python-docx live in ``ragorc[loaders]``. A
module-level import would make ``import ragorc`` fail for a deployment that only
ever ingests text, so each one is imported at call time and the ``ImportError``
names the extra to install.

Decoding is a failure mode, not a detail
----------------------------------------
Real corpora contain files that are almost UTF-8. Strict decoding is tried first;
on failure the file is decoded with replacement and flagged with
``encoding_errors`` in its metadata rather than being dropped. Losing one
character is recoverable, losing the document is not — and the flag means the
damage is visible in the store instead of being silently indexed as prose.
"""

from __future__ import annotations

import asyncio
import contextlib
import csv
import fnmatch
import re
from collections.abc import AsyncIterator, Iterable, Sequence
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Any, ClassVar, Literal

import orjson
import structlog

from ragorc.core.concurrency import bounded_gather
from ragorc.core.errors import ValidationFailed
from ragorc.core.ids import content_hash, document_id
from ragorc.core.models import Document, Modality
from ragorc.core.registry import register
from ragorc.core.settings import Settings, get_settings

log = structlog.get_logger(__name__)

__all__ = [
    "LOADERS",
    "MAX_FILE_BYTES",
    "BaseLoader",
    "CSVLoader",
    "DirectoryLoader",
    "DocxLoader",
    "HTMLLoader",
    "JSONLLoader",
    "JSONLoader",
    "MarkdownLoader",
    "PDFLoader",
    "TextLoader",
    "load",
    "loader_for",
]

MAX_FILE_BYTES = 20_000_000
"""Skip ceiling for :class:`DirectoryLoader`, matching
:class:`~ragorc.validate.schema.DocumentValidator`'s own limit. Reading a file the
validator is guaranteed to reject is pure waste, and the files that blow past it
are almost always minified bundles or data dumps rather than prose."""

_SNIFF_BYTES = 8192
"""Bytes read to decide whether a file with an unrecognized suffix is binary. A
NUL in the first 8 KiB is conclusive in practice and costs one page of I/O."""

_MAX_PATH_CHARS = 4096
"""Above this length a string is treated as a payload, not a path. ``Path.is_file``
on a 2 MB string raises ``ENAMETOOLONG``, and the guard is cheaper than the
exception."""

_CONTENT_KEYS: tuple[str, ...] = ("content", "text", "body", "page_content", "value")
"""Candidate content keys for JSON records, in priority order, used when no
``content_key`` is configured."""

_ID_KEYS: tuple[str, ...] = ("id", "_id", "uuid", "key", "slug", "url")
"""Candidate stable-key fields for JSON/CSV records. Preferred over the ordinal so
inserting a record does not re-identify the ones after it."""

_TEXT_SUFFIXES = frozenset(
    {
        ".txt", ".text", ".log", ".rst", ".org", ".tex", ".adoc", ".ini", ".cfg",
        ".conf", ".toml", ".yaml", ".yml", ".env", ".sql", ".graphql", ".proto",
    }
)  # fmt: skip

_CODE_SUFFIXES = frozenset(
    {
        ".py", ".pyi", ".js", ".jsx", ".ts", ".tsx", ".go", ".rs", ".java", ".kt",
        ".c", ".h", ".cc", ".cpp", ".hpp", ".cs", ".rb", ".php", ".swift", ".scala",
        ".sh", ".bash", ".zsh", ".ps1", ".lua", ".r", ".m", ".mm", ".dart", ".ex",
        ".exs", ".erl", ".hs", ".clj", ".vue", ".svelte",
    }
)  # fmt: skip
"""Code gets ``Modality.CODE`` so the router and the code splitter can treat it as
code without re-sniffing the content downstream."""

_BINARY_SUFFIXES = frozenset(
    {
        ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp", ".ico", ".svgz",
        ".mp3", ".mp4", ".wav", ".flac", ".avi", ".mov", ".mkv", ".webm",
        ".zip", ".gz", ".bz2", ".xz", ".7z", ".rar", ".tar", ".jar", ".war",
        ".so", ".dylib", ".dll", ".exe", ".bin", ".o", ".a", ".obj", ".class",
        ".pyc", ".pyo", ".whl", ".egg", ".pdb", ".db", ".sqlite", ".sqlite3",
        ".parquet", ".avro", ".onnx", ".pt", ".pth", ".safetensors", ".npy", ".npz",
        ".woff", ".woff2", ".ttf", ".otf", ".eot", ".psd", ".ai", ".sketch",
    }
)  # fmt: skip

_SKIP_DIRS = frozenset(
    {
        ".git", ".hg", ".svn", ".venv", "venv", "env", "node_modules", "__pycache__",
        ".mypy_cache", ".ruff_cache", ".pytest_cache", ".tox", ".idea", ".vscode",
        "dist", "build", "target", ".next", ".nuxt", ".cache", "site-packages",
    }
)  # fmt: skip
"""Directories that never contain corpus text but frequently contain hundreds of
thousands of files. Pruning them at discovery time is the difference between a
directory walk that takes a second and one that takes minutes."""

_FRONT_MATTER = re.compile(r"\A---[ \t]*\r?\n(?P<block>.*?)(?:\r?\n)?^---[ \t]*\r?\n?", re.S | re.M)
_ATX_TITLE = re.compile(r"^#\s+(?P<title>.+?)\s*#*\s*$", re.M)
_BLANK_RUN = re.compile(r"\n{3,}")
_INLINE_LIST = re.compile(r"\A\[(?P<items>.*)\]\Z", re.S)


# ---------------------------------------------------------------------------
# Reading and decoding
# ---------------------------------------------------------------------------
def _decode(raw: bytes) -> tuple[str, bool]:
    """Decode bytes to text, reporting whether characters had to be replaced.

    ``utf-8-sig`` first so a BOM never becomes a leading zero-width character in
    the content — which would shift every chunk offset by one and show up in the
    generator's context as a stray glyph.
    """
    try:
        return raw.decode("utf-8-sig"), False
    except UnicodeDecodeError:
        return raw.decode("utf-8", errors="replace"), True


def _read_file_sync(path: Path) -> tuple[str, dict[str, Any]]:
    """Read + decode + stat in one thread hop. Blocking by design."""
    raw = path.read_bytes()
    text, replaced = _decode(raw)
    stat = path.stat()
    meta: dict[str, Any] = {
        "bytes": len(raw),
        "modified_at": datetime.fromtimestamp(stat.st_mtime, tz=UTC).isoformat(),
    }
    if replaced:
        meta["encoding_errors"] = True
        log.warning("loader_encoding_errors", path=str(path))
    return text, meta


def _read_bytes_sync(path: Path) -> tuple[bytes, dict[str, Any]]:
    raw = path.read_bytes()
    stat = path.stat()
    return raw, {
        "bytes": len(raw),
        "modified_at": datetime.fromtimestamp(stat.st_mtime, tz=UTC).isoformat(),
    }


def _as_path(source: Any) -> Path | None:
    """Interpret ``source`` as a filesystem path, or ``None`` if it cannot be one."""
    if isinstance(source, Path):
        return source
    if isinstance(source, str) and source and "\n" not in source and len(source) <= _MAX_PATH_CHARS:
        with contextlib.suppress(OSError, ValueError):
            return Path(source)
    return None


def _existing_file(source: Any) -> Path | None:
    path = _as_path(source)
    if path is None:
        return None
    with contextlib.suppress(OSError):
        if path.is_file():
            return path
    return None


def _modality_for(suffix: str) -> Modality:
    if suffix in _CODE_SUFFIXES:
        return Modality.CODE
    return Modality.TEXT


def _collapse(text: str) -> str:
    """Collapse runs of blank lines. Extracted HTML and PDF text is full of them,
    and every one costs a token in the generation context for no information."""
    return _BLANK_RUN.sub("\n\n", text).strip()


def _json_safe(value: Any) -> Any:
    """Make a parsed value safe to carry in ``Document.metadata``.

    Metadata lands in a Postgres ``jsonb`` column and in a Qdrant payload, so
    anything that is not a JSON scalar or container is flattened here rather than
    left for ``orjson``'s ``default=str`` to guess at three layers down.
    """
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, dict):
        return {str(k): _json_safe(v) for k, v in value.items()}
    if isinstance(value, (list, tuple, set, frozenset)):
        return [_json_safe(v) for v in value]
    return str(value)


# ---------------------------------------------------------------------------
# Front matter
# ---------------------------------------------------------------------------
def _parse_front_matter(text: str) -> tuple[dict[str, Any], str]:
    """Split leading ``---`` front matter from the body.

    PyYAML is used when it is importable and a flat parser is used when it is not,
    because YAML is not a base dependency of this library and front matter in
    practice is a mapping of scalars and short lists. The fallback covers that
    shape exactly and refuses to invent structure for anything more complex — a
    silently mis-parsed nested block would become wrong metadata, which is worse
    than absent metadata since filters would then quietly exclude the document.
    """
    match = _FRONT_MATTER.match(text)
    if match is None:
        return {}, text
    block = match.group("block")
    body = text[match.end() :]
    try:
        import yaml
    except ImportError:
        return _flat_mapping(block), body
    try:
        parsed = yaml.safe_load(block)
    except yaml.YAMLError as exc:
        log.warning("front_matter_unparsable", error=str(exc)[:200])
        return _flat_mapping(block), body
    if not isinstance(parsed, dict):
        return _flat_mapping(block), body
    return {str(k): _json_safe(v) for k, v in parsed.items()}, body


def _flat_mapping(block: str) -> dict[str, Any]:
    """``key: value`` and ``- item`` continuation lists. Nothing else."""
    out: dict[str, Any] = {}
    current: str | None = None
    for raw_line in block.splitlines():
        line = raw_line.rstrip()
        if not line.strip() or line.lstrip().startswith("#"):
            continue
        stripped = line.strip()
        if stripped.startswith("- ") and current is not None:
            bucket = out.setdefault(current, [])
            if isinstance(bucket, list):
                bucket.append(_scalar(stripped[2:]))
            continue
        if ":" not in stripped:
            continue
        key, _, value = stripped.partition(":")
        key = key.strip()
        if not key:
            continue
        text = value.strip()
        if not text:
            # A bare "key:" opens a block list; the value arrives on the
            # following "- item" lines.
            out[key] = []
            current = key
            continue
        out[key] = _scalar(text)
        current = key
    return out


def _scalar(token: str) -> Any:
    """Coerce a front-matter scalar. Quoted stays a string, unquoted is typed."""
    text = token.strip()
    if len(text) >= 2 and text[0] == text[-1] and text[0] in "\"'":
        return text[1:-1]
    inline = _INLINE_LIST.match(text)
    if inline is not None:
        items = inline.group("items").strip()
        return [_scalar(part) for part in items.split(",")] if items else []
    lowered = text.casefold()
    if lowered in ("true", "yes"):
        return True
    if lowered in ("false", "no"):
        return False
    if lowered in ("null", "none", "~"):
        return None
    with contextlib.suppress(ValueError):
        return int(text)
    with contextlib.suppress(ValueError):
        return float(text)
    return text


# ---------------------------------------------------------------------------
# Base loader
# ---------------------------------------------------------------------------
class BaseLoader:
    """Shared identity, metadata and threading policy for every loader.

    Subclasses implement :meth:`load`. Everything that decides whether ingest is
    idempotent — the id, the checksum, the tenant, the source label — is built in
    exactly one place, :meth:`_document`, so a new loader cannot get it wrong by
    omission.
    """

    name: ClassVar[str] = "base"
    suffixes: ClassVar[tuple[str, ...]] = ()
    modality: ClassVar[Modality] = Modality.TEXT

    def __init__(
        self,
        *,
        tenant_id: str | None = None,
        metadata: dict[str, Any] | None = None,
        settings: Settings | None = None,
    ) -> None:
        self.settings = settings or get_settings()
        self.tenant_id = tenant_id or self.settings.tenant_id
        self.base_metadata = dict(metadata or {})

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        raise NotImplementedError(f"{type(self).__name__} must implement load(source)")

    async def load_many(self, sources: Iterable[Any], **kwargs: Any) -> list[Document]:
        """Load several sources concurrently, bounded by ``max_concurrent_documents``.

        Bounded rather than a bare gather: the fan-out is over a corpus, and 10k
        simultaneous ``to_thread`` hops would exhaust the default executor and hold
        10k decoded file bodies in memory at once.
        """
        items = list(sources)
        if not items:
            return []
        batches = await bounded_gather(
            (self.load(item, **kwargs) for item in items),
            limit=max(1, self.settings.indexing.max_concurrent_documents),
        )
        return [doc for batch in batches for doc in batch]

    # -- identity ---------------------------------------------------------
    def _document(
        self,
        content: str,
        *,
        source: str,
        fragment: str | None = None,
        title: str | None = None,
        metadata: dict[str, Any] | None = None,
        modality: Modality | None = None,
    ) -> Document:
        """Build a Document with a deterministic id and a mandatory checksum."""
        label = f"{source}#{fragment}" if fragment else source
        meta: dict[str, Any] = {**self.base_metadata, "loader": self.name}
        if metadata:
            meta.update(metadata)
        return Document(
            id=document_id(label, tenant_id=self.tenant_id),
            content=content,
            metadata=meta,
            source=label,
            title=title,
            modality=modality or self.modality,
            checksum=content_hash(content),
            tenant_id=self.tenant_id,
        )

    # -- source resolution ------------------------------------------------
    def _require_file(self, source: Any) -> Path:
        path = _existing_file(source)
        if path is None:
            raise ValidationFailed(
                "loader source is not a readable file",
                loader=self.name,
                source=str(source)[:200],
            )
        return path

    def _file_or_payload(self, source: Any) -> tuple[Path | None, str | None]:
        """Accept a path *or* the payload itself.

        HTML fetched over the network and JSON assembled in memory never touch the
        filesystem, and forcing a caller to write a temp file to use a loader is
        the kind of friction that gets worked around with a private copy of the
        parsing code.
        """
        if isinstance(source, (bytes, bytearray)):
            return None, _decode(bytes(source))[0]
        path = _existing_file(source)
        if path is not None:
            return path, None
        if isinstance(source, str) and source.strip():
            return None, source
        raise ValidationFailed(
            "loader source is neither a readable file nor a payload",
            loader=self.name,
            source=str(source)[:200],
        )

    async def _read(self, source: Any) -> tuple[str, str, dict[str, Any]]:
        """Return ``(text, source_label, file_metadata)`` for a path or payload."""
        path, payload = self._file_or_payload(source)
        if path is not None:
            text, meta = await asyncio.to_thread(_read_file_sync, path)
            return text, str(path), meta
        assert payload is not None  # noqa: S101 - narrowing; _file_or_payload raises otherwise
        return payload, f"inline:{content_hash(payload, size=8)}", {"bytes": len(payload.encode())}


# ---------------------------------------------------------------------------
# Plain text and markdown
# ---------------------------------------------------------------------------
@register("loader", "text")
class TextLoader(BaseLoader):
    """UTF-8 text, one document per file."""

    name = "text"
    suffixes = (".txt", ".text", ".log", ".rst", ".org", ".tex", ".adoc")

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        path = self._require_file(source)
        text, meta = await asyncio.to_thread(_read_file_sync, path)
        if not text.strip():
            log.warning("loader_empty_file", loader=self.name, path=str(path))
            return []
        meta["extension"] = path.suffix.lower()
        return [
            self._document(
                text,
                source=str(path),
                title=path.stem,
                metadata=meta,
                modality=_modality_for(path.suffix.lower()),
            )
        ]


@register("loader", "markdown", "md")
class MarkdownLoader(BaseLoader):
    """Markdown with YAML front matter folded into metadata.

    The front matter is *removed* from the content rather than left in place.
    Indexed as prose it is noise that matches every query about dates and authors;
    lifted into metadata it becomes filterable, which is what it is for. The title
    is taken from the front matter when present and from the first ATX heading
    otherwise, because ``title`` is what the context packer shows the model as the
    provenance of a chunk.
    """

    name = "markdown"
    suffixes = (".md", ".markdown", ".mdown", ".mkd", ".mdx")

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        path = self._require_file(source)
        text, meta = await asyncio.to_thread(_read_file_sync, path)
        front, body = await asyncio.to_thread(_parse_front_matter, text)
        body = body.lstrip("\n")
        if not body.strip():
            log.warning("loader_empty_file", loader=self.name, path=str(path))
            return []

        title = front.get("title")
        if not isinstance(title, str) or not title.strip():
            heading = _ATX_TITLE.search(body)
            title = heading.group("title") if heading else path.stem

        meta.update(front)
        meta["extension"] = path.suffix.lower()
        if front:
            # Recorded so a retrieval filter can distinguish authored metadata
            # from metadata this loader inferred.
            meta["front_matter_keys"] = sorted(front)
        return [self._document(body, source=str(path), title=str(title), metadata=meta)]


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------
def _pick(record: dict[str, Any], configured: str | None, candidates: Sequence[str]) -> str | None:
    if configured is not None:
        return configured if configured in record else None
    return next((key for key in candidates if key in record), None)


class _JSONBase(BaseLoader):
    """Shared record-to-document mapping for the JSON and JSONL loaders."""

    def __init__(
        self,
        *,
        content_key: str | None = None,
        id_key: str | None = None,
        tenant_id: str | None = None,
        metadata: dict[str, Any] | None = None,
        settings: Settings | None = None,
    ) -> None:
        super().__init__(tenant_id=tenant_id, metadata=metadata, settings=settings)
        self.content_key = content_key
        self.id_key = id_key

    def _record_to_document(
        self,
        record: Any,
        *,
        source: str,
        ordinal: int,
        file_meta: dict[str, Any],
    ) -> Document | None:
        """Map one record. ``None`` when it carries no usable text."""
        if isinstance(record, str):
            return (
                self._document(
                    record, source=source, fragment=f"item={ordinal}", metadata=dict(file_meta)
                )
                if record.strip()
                else None
            )
        if not isinstance(record, dict):
            log.warning("json_record_skipped", source=source, ordinal=ordinal, reason="not_object")
            return None

        key = _pick(record, self.content_key, _CONTENT_KEYS)
        if key is None:
            log.warning(
                "json_record_skipped",
                source=source,
                ordinal=ordinal,
                reason="no_content_key",
                configured=self.content_key,
                available=sorted(record)[:10],
            )
            return None
        content = record[key]
        if not isinstance(content, str):
            content = "" if content is None else str(content)
        if not content.strip():
            return None

        # Everything that is not the content becomes metadata: a JSON corpus
        # carries its filters in the same object as its text, and dropping the
        # siblings would throw away the only structured signal the source has.
        meta = {**file_meta, **{k: _json_safe(v) for k, v in record.items() if k != key}}
        id_field = _pick(record, self.id_key, _ID_KEYS)
        stable = record.get(id_field) if id_field else None
        fragment = f"{id_field}={stable}" if stable not in (None, "") else f"item={ordinal}"
        title = record.get("title") or record.get("name")
        return self._document(
            content,
            source=source,
            fragment=fragment,
            title=str(title) if isinstance(title, (str, int, float)) else None,
            metadata=meta,
        )


@register("loader", "json")
class JSONLoader(_JSONBase):
    """A JSON array (or single object) of records.

    ``content_key`` names the field holding the text; without it the first of
    ``content``/``text``/``body``/``page_content``/``value`` present wins. Every
    other field is folded into metadata.
    """

    name = "json"
    suffixes = (".json",)

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        text, label, file_meta = await self._read(source)
        data = await asyncio.to_thread(self._parse, text, label)
        records = data if isinstance(data, list) else [data]
        docs: list[Document] = []
        for ordinal, record in enumerate(records):
            doc = self._record_to_document(
                record, source=label, ordinal=ordinal, file_meta=file_meta
            )
            if doc is not None:
                docs.append(doc)
        log.info("json_loaded", source=label, records=len(records), documents=len(docs))
        return docs

    @staticmethod
    def _parse(text: str, label: str) -> Any:
        try:
            return orjson.loads(text)
        except orjson.JSONDecodeError as exc:
            raise ValidationFailed(
                "source is not valid JSON", source=label, error=str(exc)
            ) from exc


@register("loader", "jsonl", "ndjson")
class JSONLLoader(_JSONBase):
    """Newline-delimited JSON: one record per line.

    A malformed line is warned about and skipped rather than fatal. JSONL is the
    export format of log pipelines and scrapers, and a single truncated line at
    the tail of a 2 GB file is the normal case, not an exceptional one — aborting
    would throw away every record that parsed.
    """

    name = "jsonl"
    suffixes = (".jsonl", ".ndjson")

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        text, label, file_meta = await self._read(source)
        return await asyncio.to_thread(self._build, text, label, file_meta)

    def _build(self, text: str, label: str, file_meta: dict[str, Any]) -> list[Document]:
        docs: list[Document] = []
        bad = 0
        for ordinal, line in enumerate(text.splitlines()):
            if not line.strip():
                continue
            record = _loads_or_none(line)
            if record is None:
                bad += 1
                continue
            doc = self._record_to_document(
                record, source=label, ordinal=ordinal, file_meta=file_meta
            )
            if doc is not None:
                docs.append(doc)
        if bad:
            log.warning("jsonl_lines_skipped", source=label, malformed=bad, loaded=len(docs))
        return docs


def _loads_or_none(line: str) -> Any:
    """Parse one JSONL line. The try lives here so the caller's loop stays clean
    of per-iteration exception handling."""
    try:
        return orjson.loads(line)
    except orjson.JSONDecodeError:
        return None


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------
@register("loader", "csv", "tsv")
class CSVLoader(BaseLoader):
    """Delimited data, one document per row or one per file.

    Per-row is the default because a row is the unit a question is usually about,
    and a 50k-row file as one document would be split by character count into
    chunks that cut through the middle of records. Per-file exists for small
    reference tables, where the answer needs the whole table to be comparable.

    The delimiter is sniffed rather than assumed: the difference between a comma
    and a semicolon file is one column of garbage per row, and it fails silently.
    """

    name = "csv"
    suffixes = (".csv", ".tsv")
    modality = Modality.TABLE

    def __init__(
        self,
        *,
        mode: Literal["row", "file"] = "row",
        content_column: str | None = None,
        id_column: str | None = None,
        delimiter: str | None = None,
        tenant_id: str | None = None,
        metadata: dict[str, Any] | None = None,
        settings: Settings | None = None,
    ) -> None:
        super().__init__(tenant_id=tenant_id, metadata=metadata, settings=settings)
        self.mode = mode
        self.content_column = content_column
        self.id_column = id_column
        self.delimiter = delimiter

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        path = self._require_file(source)
        text, meta = await asyncio.to_thread(_read_file_sync, path)
        if not text.strip():
            return []
        meta["extension"] = path.suffix.lower()
        if self.mode == "file":
            return [
                self._document(_collapse(text), source=str(path), title=path.stem, metadata=meta)
            ]
        return await asyncio.to_thread(self._rows, text, str(path), path.stem, meta)

    def _rows(self, text: str, label: str, stem: str, file_meta: dict[str, Any]) -> list[Document]:
        delimiter = self.delimiter or _sniff_delimiter(text, label)
        reader = csv.DictReader(text.splitlines(), delimiter=delimiter)
        docs: list[Document] = []
        for ordinal, row in enumerate(reader, start=1):
            clean = {
                str(k).strip(): ("" if v is None else v)
                for k, v in row.items()
                if k is not None and str(k).strip()
            }
            if not clean:
                continue
            content = self._row_content(clean, label, ordinal)
            if content is None or not content.strip():
                continue
            stable = clean.get(self.id_column) if self.id_column else None
            if stable in (None, ""):
                stable = next((clean[k] for k in _ID_KEYS if clean.get(k)), None)
                fragment = f"row={ordinal}" if stable in (None, "") else f"id={stable}"
            else:
                fragment = f"{self.id_column}={stable}"
            meta: dict[str, Any] = {**file_meta, "row": ordinal, "delimiter": delimiter}
            if self.content_column:
                meta.update(
                    {k: _json_safe(v) for k, v in clean.items() if k != self.content_column}
                )
            docs.append(
                self._document(
                    content,
                    source=label,
                    fragment=fragment,
                    title=f"{stem} row {ordinal}",
                    metadata=meta,
                )
            )
        log.info("csv_loaded", source=label, documents=len(docs), delimiter=delimiter)
        return docs

    def _row_content(self, row: dict[str, str], label: str, ordinal: int) -> str | None:
        if self.content_column is not None:
            value = row.get(self.content_column)
            if value is None:
                log.warning(
                    "csv_row_skipped",
                    source=label,
                    row=ordinal,
                    reason="missing_content_column",
                    column=self.content_column,
                )
                return None
            return value
        # No designated text column: render the row as labelled lines. Keeping the
        # column names is what makes the row retrievable — "revenue" is in the
        # header, not in the cell, and a bare list of values matches nothing.
        return "\n".join(f"{key}: {value}" for key, value in row.items() if str(value).strip())


def _sniff_delimiter(text: str, label: str) -> str:
    """Detect the delimiter from the header plus a few rows."""
    sample = text[:8192]
    try:
        return csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
    except csv.Error:
        # Sniffing fails on single-column files, which are unambiguous anyway.
        guess = "\t" if "\t" in sample.splitlines()[0] else ","
        log.debug("csv_delimiter_guessed", source=label, delimiter=guess)
        return guess


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------
@register("loader", "html", "htm")
class HTMLLoader(BaseLoader):
    """HTML via BeautifulSoup + lxml, reduced to its main content.

    Two reductions, both of which change retrieval quality rather than tidiness:

    * **Chrome is removed** (``script``, ``style``, ``nav``, ``header``,
      ``footer``, ``aside``, ``form``). Site chrome repeats on every page of a
      corpus, so it is the most duplicated text in the index; it matches queries
      about menu labels and pushes real content out of the candidate window.
    * **The main region is preferred** over the whole body. A 200-word article
      inside a page with 2000 words of navigation and related-links would
      otherwise be a fifth of its own document, and the semantic splitter's
      percentile threshold is computed over all of it.

    lxml is the parser because it is several times faster than ``html.parser`` and
    recovers from real-world broken markup instead of truncating at it; the
    stdlib parser is the fallback when lxml is not installed.
    """

    name = "html"
    suffixes = (".html", ".htm", ".xhtml")

    _STRIP: ClassVar[tuple[str, ...]] = (
        "script", "style", "nav", "header", "footer", "aside", "form", "noscript",
        "iframe", "svg", "template", "button",
    )  # fmt: skip

    _MAIN_SELECTORS: ClassVar[tuple[str, ...]] = (
        "main",
        "article",
        "[role=main]",
        "#main",
        "#content",
        "#main-content",
        ".main-content",
        ".post-content",
        ".entry-content",
        ".markdown-body",
        ".content",
    )

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        text, label, file_meta = await self._read(source)
        content, meta = await asyncio.to_thread(self._extract, text)
        if not content:
            log.warning("loader_empty_file", loader=self.name, path=label)
            return []
        file_meta.update(meta)
        return [
            self._document(
                content,
                source=label,
                title=meta.get("title") or Path(label).stem,
                metadata=file_meta,
            )
        ]

    def _extract(self, html: str) -> tuple[str, dict[str, Any]]:
        try:
            from bs4 import BeautifulSoup
        except ImportError as exc:  # pragma: no cover - exercised by the extra
            raise ImportError(
                "HTMLLoader needs beautifulsoup4 and lxml: pip install 'ragorc[loaders]'"
            ) from exc

        try:
            soup = BeautifulSoup(html, "lxml")
            parser = "lxml"
        except Exception:  # bs4 raises FeatureNotFound, which is not an ImportError
            soup = BeautifulSoup(html, "html.parser")
            parser = "html.parser"
            log.warning(
                "html_parser_fallback",
                parser=parser,
                hint="pip install 'ragorc[loaders]' for lxml",
            )

        meta: dict[str, Any] = {"parser": parser}
        if soup.title and soup.title.string:
            meta["title"] = soup.title.string.strip()
        for tag in soup.find_all("meta"):
            key = tag.get("name") or tag.get("property")
            value = tag.get("content")
            if key in ("description", "og:title", "og:description", "author") and value:
                meta[str(key).replace("og:", "")] = str(value).strip()
        canonical = soup.find("link", rel="canonical")
        if canonical is not None and canonical.get("href"):
            meta["canonical_url"] = str(canonical["href"])
        if soup.html is not None and soup.html.get("lang"):
            meta["language"] = str(soup.html["lang"])

        for tag in soup(list(self._STRIP)):
            tag.decompose()

        # Candidates are tried in priority order rather than as one comma-joined
        # selector: `select_one` with a list returns the first match in *document*
        # order, which would prefer an early sidebar `.content` over the `article`
        # that follows it.
        region = next(
            (found for sel in self._MAIN_SELECTORS if (found := soup.select_one(sel)) is not None),
            None,
        )
        if region is None:
            region = soup.body or soup
        else:
            meta["main_selector"] = next(
                sel for sel in self._MAIN_SELECTORS if soup.select_one(sel) is region
            )

        if not meta.get("title"):
            h1 = region.find("h1")
            if h1 is not None:
                meta["title"] = h1.get_text(strip=True)
        return _collapse(region.get_text("\n", strip=True)), meta


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
@register("loader", "pdf")
class PDFLoader(BaseLoader):
    """PDF text via PyMuPDF, per page or as one document.

    ``mode="page"`` gives one Document per page: page-level provenance survives
    into every citation, which is what a reader needs to verify an answer against
    the original. ``mode="document"`` keeps the pages joined so a paragraph that
    runs across a page break is not split by the loader before the splitter has
    seen it — and records ``page_breaks`` as ``[char_offset, page_number]`` pairs,
    so a chunk's ``start_char`` can still be resolved back to a page.

    PyMuPDF rather than pypdf or pdfminer: it is a C library with Python bindings
    and extracts text roughly an order of magnitude faster, which matters because
    a PDF corpus is the case where extraction, not embedding, dominates ingest.
    """

    name = "pdf"
    suffixes = (".pdf",)

    def __init__(
        self,
        *,
        mode: Literal["page", "document"] = "document",
        password: str | None = None,
        tenant_id: str | None = None,
        metadata: dict[str, Any] | None = None,
        settings: Settings | None = None,
    ) -> None:
        super().__init__(tenant_id=tenant_id, metadata=metadata, settings=settings)
        self.mode = mode
        self.password = password

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        path = self._require_file(source)
        # One hop for open + extract + close: the whole thing is blocking C code,
        # and splitting it would only add scheduling latency between the parts.
        pages, doc_meta = await asyncio.to_thread(self._extract, path)
        if not pages:
            log.warning("loader_empty_file", loader=self.name, path=str(path))
            return []

        stat_meta = {"pages": len(pages), **doc_meta}
        title = doc_meta.get("pdf_title") or path.stem
        if self.mode == "page":
            docs = [
                self._document(
                    text,
                    source=str(path),
                    fragment=f"page={number}",
                    title=f"{title} p.{number}",
                    metadata={**stat_meta, "page": number},
                )
                for number, text in pages
                if text.strip()
            ]
            log.info("pdf_loaded", path=str(path), pages=len(pages), documents=len(docs))
            return docs

        # Each page is normalized *before* the offsets are taken, and the joined
        # text is not touched afterwards. Collapsing whitespace after measuring
        # would shift every offset past the first collapsed run — the same class
        # of silent, unlogged mis-alignment that makes late chunking pool the
        # wrong tokens, except here it would mis-report the page of a citation.
        parts = [_collapse(text) for _, text in pages]
        breaks: list[list[int]] = []
        cursor = 0
        kept: list[str] = []
        for (number, _), part in zip(pages, parts, strict=True):
            if not part:
                continue
            breaks.append([cursor, number])
            kept.append(part)
            cursor += len(part) + 2  # the "\n\n" join below
        if not kept:
            return []
        return [
            self._document(
                "\n\n".join(kept),
                source=str(path),
                title=str(title),
                metadata={**stat_meta, "page_breaks": breaks},
            )
        ]

    def _extract(self, path: Path) -> tuple[list[tuple[int, str]], dict[str, Any]]:
        try:
            import pymupdf
        except ImportError:
            try:
                import fitz  # PyMuPDF < 1.24 shipped as `fitz`
            except ImportError as exc:
                raise ImportError("PDFLoader needs PyMuPDF: pip install 'ragorc[loaders]'") from exc
            pymupdf = fitz

        document = pymupdf.open(path)
        try:
            if document.needs_pass and not document.authenticate(self.password or ""):
                raise ValidationFailed(
                    "PDF is encrypted and the password was rejected", path=str(path)
                )
            raw = document.metadata or {}
            meta = {
                f"pdf_{key}": str(value).strip()
                for key, value in raw.items()
                if key in ("title", "author", "subject", "creator") and value
            }
            # `iter()` is not decoration: pymupdf's Document is iterable only
            # through the legacy `__getitem__` protocol, which a type checker
            # cannot see. This is the same iterator `for page in document` builds.
            pages = [
                (index + 1, page.get_text("text")) for index, page in enumerate(iter(document))
            ]
        finally:
            document.close()
        return pages, meta


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
@register("loader", "docx")
class DocxLoader(BaseLoader):
    """Word documents via python-docx, in authored order.

    The body's XML children are walked rather than reading ``doc.paragraphs`` and
    ``doc.tables`` separately, because those two collections lose the interleaving:
    a table read out of order arrives after 40 pages of prose, detached from the
    sentence that introduced it, and a chunk of bare numbers with no caption is
    unretrievable and uninterpretable.
    """

    name = "docx"
    suffixes = (".docx",)

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        path = self._require_file(source)
        content, meta = await asyncio.to_thread(self._extract, path)
        if not content:
            log.warning("loader_empty_file", loader=self.name, path=str(path))
            return []
        return [
            self._document(
                content,
                source=str(path),
                title=meta.get("docx_title") or path.stem,
                metadata=meta,
            )
        ]

    def _extract(self, path: Path) -> tuple[str, dict[str, Any]]:
        try:
            import docx
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except ImportError as exc:
            raise ImportError(
                "DocxLoader needs python-docx: pip install 'ragorc[loaders]'"
            ) from exc

        document = docx.Document(str(path))
        blocks: list[str] = []
        tables = 0
        paragraphs = 0
        for child in document.element.body.iterchildren():
            tag = str(child.tag)
            if tag.endswith("}p"):
                text = Paragraph(child, document).text.strip()
                if text:
                    blocks.append(text)
                    paragraphs += 1
            elif tag.endswith("}tbl"):
                rendered = _render_table(Table(child, document))
                if rendered:
                    blocks.append(rendered)
                    tables += 1

        meta: dict[str, Any] = {"paragraphs": paragraphs, "tables": tables}
        core = document.core_properties
        for field in ("title", "author", "subject"):
            value = getattr(core, field, None)
            if isinstance(value, str) and value.strip():
                meta[f"docx_{field}"] = value.strip()
        return _collapse("\n\n".join(blocks)), meta


def _render_table(table: Any) -> str:
    """Pipe-delimited rows. A markdown-ish table survives chunking legibly and
    keeps the header adjacent to its values, which a bare cell dump does not."""
    rows = [
        " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells) for row in table.rows
    ]
    return "\n".join(row for row in rows if row.replace("|", "").strip())


# ---------------------------------------------------------------------------
# Suffix dispatch
# ---------------------------------------------------------------------------
LOADERS: dict[str, type[BaseLoader]] = {
    suffix: cls
    for cls in (
        TextLoader,
        MarkdownLoader,
        JSONLoader,
        JSONLLoader,
        CSVLoader,
        HTMLLoader,
        PDFLoader,
        DocxLoader,
    )
    for suffix in cls.suffixes
}
"""Suffix -> loader. The registry (``resolve("loader", name)``) maps *names* for
configuration; this maps extensions for dispatch. Both point at the same classes,
so a third-party loader registered under a name can be added here by suffix
without touching the dispatcher."""


def loader_for(suffix: str) -> type[BaseLoader]:
    """Loader class for a file suffix, defaulting to :class:`TextLoader`.

    Defaulting to text rather than raising is deliberate: a corpus of ``.py``,
    ``.sql`` and ``.conf`` files is text with an unfamiliar extension, and the
    binary formats that must *not* be read this way are excluded by
    :data:`_BINARY_SUFFIXES` at discovery time instead.
    """
    return LOADERS.get(suffix.lower(), TextLoader)


# ---------------------------------------------------------------------------
# Directory
# ---------------------------------------------------------------------------
@register("loader", "directory", "dir")
class DirectoryLoader(BaseLoader):
    """Recursive directory walk, dispatching each file by extension.

    **One unreadable file must not abort a 10k-file ingest.** A corpus contains a
    truncated PDF, a file the process cannot read, a ``.docx`` that is actually a
    renamed zip — and discovering that after eight hours of embedding, with
    nothing written, is the failure mode that makes people stop trusting an
    ingest tool. So every per-file failure is caught, counted, logged with its
    path, and recorded in :attr:`failures`; the walk continues. The one thing that
    *does* abort is a bad root, because that is a configuration error and every
    subsequent file would fail the same way.

    Discovery is deliberately deterministic: ``rglob`` order depends on the
    filesystem, and ingest order decides batch composition, so the paths are
    sorted before loading. Two runs over the same tree then produce the same
    batches and the same logs.

    Cheap exclusions happen before any read: pruned directories, the binary
    suffix list, and the size ceiling. Only files with an *unrecognized* suffix
    pay a NUL-byte sniff of their first 8 KiB — checking every file would double
    the syscalls for a corpus that is mostly ``.md``.
    """

    name = "directory"

    def __init__(
        self,
        *,
        include: Sequence[str] | None = None,
        exclude: Sequence[str] | None = None,
        recursive: bool = True,
        max_bytes: int = MAX_FILE_BYTES,
        limit: int | None = None,
        concurrency: int | None = None,
        loader_kwargs: dict[str, dict[str, Any]] | None = None,
        tenant_id: str | None = None,
        metadata: dict[str, Any] | None = None,
        settings: Settings | None = None,
    ) -> None:
        super().__init__(tenant_id=tenant_id, metadata=metadata, settings=settings)
        self.include = tuple(include or ())
        self.exclude = tuple(exclude or ())
        self.recursive = recursive
        self.max_bytes = max_bytes
        self.limit = limit
        self.concurrency = concurrency or max(1, self.settings.indexing.max_concurrent_documents)
        self.loader_kwargs = dict(loader_kwargs or {})
        self.failures: list[tuple[str, str]] = []
        self.skipped: int = 0

    async def load(self, source: Any, **kwargs: Any) -> list[Document]:
        root = _as_path(source)
        if root is None or not root.is_dir():
            raise ValidationFailed("DirectoryLoader source is not a directory", source=str(source))

        self.failures = []
        self.skipped = 0
        files = await asyncio.to_thread(self._discover, root)
        if not files:
            log.warning("directory_empty", root=str(root), include=self.include)
            return []

        batches = await bounded_gather(
            (self._load_one(path, root) for path in files), limit=self.concurrency
        )
        docs = [doc for batch in batches for doc in batch]
        log.info(
            "directory_loaded",
            root=str(root),
            files=len(files),
            documents=len(docs),
            skipped=self.skipped,
            failed=len(self.failures),
        )
        return docs

    async def iter_documents(self, root: Path, *, window: int) -> AsyncIterator[list[Document]]:
        """Yield loaded documents in windows of ``window`` files.

        Discovery is cheap — it is a walk producing paths, not content — so the
        expensive part can be bounded. :meth:`load` materializes every document
        first, which for a large corpus means the whole corpus's *text* is resident
        before the first vector is written; the chunk stream is bounded but the
        document list is not, and at 100k documents the document list is the larger
        number. Windowing lets the caller hold one window at a time.
        """
        root = _as_path(root) or root
        if not root.is_dir():
            raise ValidationFailed("DirectoryLoader source is not a directory", source=str(root))
        self.failures = []
        self.skipped = 0
        paths = await asyncio.to_thread(self._discover, root)
        step = max(1, window)
        for start in range(0, len(paths), step):
            batches = await bounded_gather(
                (self._load_one(path, root) for path in paths[start : start + step]),
                limit=self.concurrency,
            )
            yield [doc for batch in batches for doc in batch]
        log.info(
            "directory_streamed",
            root=str(root),
            files=len(paths),
            window=step,
            skipped=self.skipped,
            failed=len(self.failures),
        )

    # -- discovery --------------------------------------------------------
    def _discover(self, root: Path) -> list[Path]:
        """Walk the tree and apply every cheap filter. Blocking; runs in a thread."""
        found: list[Path] = []
        for path in sorted(root.rglob("*") if self.recursive else root.glob("*")):
            if self.limit is not None and len(found) >= self.limit:
                break
            if not self._admit(path, root):
                continue
            found.append(path)
        return found

    def _admit(self, path: Path, root: Path) -> bool:
        if path.is_dir() or path.is_symlink():
            return False
        if any(part in _SKIP_DIRS for part in path.parts):
            return False
        try:
            relative = path.relative_to(root).as_posix()
        except ValueError:  # pragma: no cover - rglob results are always under root
            relative = path.name
        if self.exclude and self._matches(relative, path.name, self.exclude):
            return False
        if self.include and not self._matches(relative, path.name, self.include):
            return False

        suffix = path.suffix.lower()
        if suffix in _BINARY_SUFFIXES:
            self.skipped += 1
            return False
        try:
            size = path.stat().st_size
        except OSError as exc:
            self._record(path, exc)
            return False
        if size == 0:
            self.skipped += 1
            return False
        if size > self.max_bytes:
            self.skipped += 1
            log.warning("file_too_large", path=str(path), bytes=size, limit=self.max_bytes)
            return False
        if suffix not in LOADERS and suffix not in _TEXT_SUFFIXES and suffix not in _CODE_SUFFIXES:
            return not self._looks_binary(path)
        return True

    @staticmethod
    def _matches(relative: str, name: str, patterns: Sequence[str]) -> bool:
        """Match a pattern against the relative path *and* the bare filename, so
        ``*.md`` means "markdown anywhere" rather than only at the root."""
        return any(
            fnmatch.fnmatch(relative, pattern) or fnmatch.fnmatch(name, pattern)
            for pattern in patterns
        )

    def _looks_binary(self, path: Path) -> bool:
        try:
            with path.open("rb") as handle:
                head = handle.read(_SNIFF_BYTES)
        except OSError as exc:
            self._record(path, exc)
            return True
        if b"\x00" in head:
            self.skipped += 1
            return True
        return False

    # -- loading ----------------------------------------------------------
    async def _load_one(self, path: Path, root: Path) -> list[Document]:
        """Load one file, converting any failure into a recorded warning."""
        cls = loader_for(path.suffix)
        try:
            loader = cls(
                tenant_id=self.tenant_id,
                metadata={**self.base_metadata, "root": str(root)},
                settings=self.settings,
                **self.loader_kwargs.get(cls.name, {}),
            )
            return await loader.load(path)
        except (OSError, ValueError, TypeError, ImportError, ValidationFailed) as exc:
            self._record(path, exc)
            return []

    def _record(self, path: Path, exc: BaseException) -> None:
        self.failures.append((str(path), f"{type(exc).__name__}: {exc}"))
        log.warning(
            "file_load_failed",
            path=str(path),
            error=str(exc)[:200],
            error_type=type(exc).__name__,
        )


# ---------------------------------------------------------------------------
# The dispatcher
# ---------------------------------------------------------------------------
async def load(
    source: Any,
    *,
    tenant_id: str | None = None,
    metadata: dict[str, Any] | None = None,
    settings: Settings | None = None,
    **kwargs: Any,
) -> list[Document]:
    """Load anything: a file, a directory, or an iterable of either.

    Dispatch is by suffix for files (:func:`loader_for`) and by ``is_dir`` for
    directories. This is the entry point :meth:`IngestPipeline.ingest
    <ragorc.index.pipeline.IngestPipeline.ingest>` uses when handed a path rather
    than documents, which is why it accepts a collection: ingesting "these forty
    files" is as common as ingesting a tree, and the alternative is forty awaits
    with no concurrency bound between them.
    """
    resolved = settings or get_settings()
    common: dict[str, Any] = {"tenant_id": tenant_id, "metadata": metadata, "settings": resolved}

    path = _as_path(source)
    if path is not None and path.is_dir():
        return await DirectoryLoader(**common, **kwargs).load(path)
    if path is not None and path.is_file():
        # The ceiling applies to a named file too, not only to files found by
        # walking a directory. `DirectoryLoader` skipped oversize files while this
        # branch read them whole into memory, so the limit was avoidable by naming
        # the file — and the validator rejects it afterwards regardless, making the
        # read pure waste. Raised rather than skipped: the caller named this path,
        # and a silent skip is an ingest that reports success having indexed
        # nothing.
        # Popped, not read: `max_bytes` is a DirectoryLoader parameter, and
        # forwarding it to a single-file loader is a TypeError. It was already one
        # before this check existed, which is a trap either way — the same call
        # works on a directory and fails on a file.
        ceiling = int(kwargs.pop("max_bytes", None) or MAX_FILE_BYTES)
        size = path.stat().st_size
        if size > ceiling:
            raise ValidationFailed(
                "file exceeds the loader size ceiling",
                path=str(path),
                bytes=size,
                limit_bytes=ceiling,
            )
        return await loader_for(path.suffix)(**common, **kwargs).load(path)

    if isinstance(source, (str, bytes, bytearray)):
        # A payload rather than a path: only the structured loaders can consume
        # one, and JSON is the only shape we can identify without guessing.
        return await JSONLoader(**common, **kwargs).load(source)

    if isinstance(source, Iterable):
        items = list(source)
        batches = await bounded_gather(
            (load(item, **common, **kwargs) for item in items),
            limit=max(1, resolved.indexing.max_concurrent_documents),
        )
        return [doc for batch in batches for doc in batch]

    raise ValidationFailed("unsupported ingest source", source=str(source)[:200])
